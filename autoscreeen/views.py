import time
import threading
import logging
import os
from django.shortcuts import render
from django.http import HttpResponse
from django.conf import settings
from docx import Document
from docx.shared import Inches
import xlsxwriter
import pyautogui
import keyboard
from datetime import datetime

# Set up logging
logging.basicConfig(level=logging.DEBUG, format='%(asctime)s - %(levelname)s - %(message)s')

# Global variables
screenshot_mode = False
screenshot_thread = None
file_path = None
file_format = None
workbook = None
worksheet = None
row_index = 0
doc = None
error_message = None
sheet_count = 0  # Counter for the number of sheets created in Excel
screenshot_folder = os.path.join(settings.MEDIA_ROOT, 'screenshots')

# Ensure the media directory and screenshots folder exist
os.makedirs(screenshot_folder, exist_ok=True)

# Thread lock for safe access to global variables
thread_lock = threading.Lock()

# Function to handle screenshots
def take_screenshot():
    global workbook, worksheet, row_index, doc
    try:
        # Capture screenshot in memory
        screenshot = pyautogui.screenshot()
        
        # Create a unique filename based on the current time
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S_%f")
        unique_screenshot_path = os.path.join(screenshot_folder, f"screenshot_{timestamp}.png")
        
        # Save the screenshot to a unique file in the screenshots folder
        screenshot.save(unique_screenshot_path, format='PNG')

        with thread_lock:
            if worksheet:
                # Insert screenshot into Excel file at the next available row
                worksheet.insert_image(row_index, 0, unique_screenshot_path)
                row_index += 65  # Move down for the next screenshot
                logging.debug(f"Inserted image into Excel at row: {row_index - 20}")

            elif doc:
                # Insert screenshot into Word document
                doc.add_picture(unique_screenshot_path, width=Inches(6))
                doc.add_paragraph()
                doc.save(file_path)
                logging.debug("Inserted image into Word document")

    except Exception as e:
        logging.error(f"Error while taking screenshot: {e}")

def listen_for_screenshots():
    global screenshot_mode
    logging.debug("Started screenshot listener thread.")
    while screenshot_mode:
        if keyboard.is_pressed('right'):
            take_screenshot()
            time.sleep(0.5)
        time.sleep(0.1)

# Home view to handle start/stop of screenshot mode and creating a new file
def home(request):
    global screenshot_mode, screenshot_thread, workbook, worksheet, row_index, file_path, file_format, error_message, sheet_count

    message = 'Screenshot mode is OFF'
    message_color = 'green'

    if error_message:
        message = error_message
        message_color = 'red'
        error_message = None

    try:
        if request.method == "POST":
            # Check if "New File" was clicked
            if 'new_file' in request.POST:
                custom_file_name = request.POST.get('file_name', '').strip()
                selected_format = request.POST.get('file_format')
                logging.debug(f"New file request: Name='{custom_file_name}', Format='{selected_format}'")

                if custom_file_name:
                    # Close any existing workbook or document and reset variables
                    if workbook:
                        workbook.close()
                    workbook = None
                    worksheet = None
                    doc = None
                    file_path = None
                    file_format = None
                    sheet_count = 0  # Reset sheet count

                    # Create a new file based on selected format
                    if selected_format == 'Word':
                        create_word_file(custom_file_name)
                    elif selected_format == 'Excel':
                        create_excel_file(custom_file_name)
                    file_format = selected_format
                    message = f'New file "{custom_file_name}.{file_format.lower()}" created.'
                    logging.debug(message)
                else:
                    message = 'Invalid file name.'
                    message_color = 'red'
                    logging.warning(message)

            # Check if "Screenshot On" was clicked
            elif 'screenshot_on' in request.POST:
                with thread_lock:
                    if not screenshot_mode and file_path:
                        screenshot_mode = True
                        if file_format == 'Excel':
                            # Create a new worksheet for screenshots if workbook already exists
                            if workbook:
                                # Create a new sheet for each new "Mode On" session
                                sheet_count += 1
                                worksheet = workbook.add_worksheet(f'Sheet_{sheet_count}')
                                row_index = 0  # Reset row index for new screenshots
                                logging.debug(f"New worksheet created in Excel file for screenshots.")
                            else:
                                create_excel_file(request.POST.get('file_name', 'default_file'))
                                logging.debug("Excel file created for screenshot mode.")

                        screenshot_thread = threading.Thread(target=listen_for_screenshots)
                        screenshot_thread.start()
                        message = 'Screenshot mode ON. Press "right" to take screenshots one by one.'
                        logging.debug(message)
                    elif screenshot_mode:
                        message = 'Screenshot mode is already ON.'

            # Check if "Screenshot Off" was clicked
            elif 'screenshot_off' in request.POST:
                with thread_lock:
                    screenshot_mode = False
                    if screenshot_thread:
                        screenshot_thread.join()

                    message = 'Screenshot mode OFF'
                    logging.debug(message)

    except Exception as e:
        message = f'An error occurred: {e}'
        message_color = 'red'
        logging.error(message)

    return render(request, 'index.html', {'message': message, 'message_color': message_color, 'file_created': bool(file_path)})

# Helper functions to create new files
def create_word_file(file_name):
    global file_path, doc
    file_path = os.path.join(settings.MEDIA_ROOT, f"{file_name}.docx")
    doc = Document()
    doc.save(file_path)
    logging.debug(f"Word file created: {file_path}")

def create_excel_file(file_name):
    global file_path, workbook, worksheet, row_index, sheet_count
    file_path = os.path.join(settings.MEDIA_ROOT, f"{file_name}.xlsx")
    workbook = xlsxwriter.Workbook(file_path)
    worksheet = workbook.add_worksheet('Screenshots')  # Initial worksheet
    row_index = 0
    sheet_count = 1  # Reset sheet count for new Excel file
    logging.debug(f"Excel file created: {file_path}")

# Function to download the last created file
def download(request):
    global file_path, file_format, workbook
    try:
        # Ensure workbook is closed before download
        if file_format == 'Excel' and workbook:
            workbook.close()
            workbook = None  # Reset to ensure no further changes are made

        # Proceed with download if file exists
        if file_path and os.path.exists(file_path):
            content_type = (
                'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                if file_format == 'Word' else
                'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
            )
            logging.debug(f"Preparing to download file: {file_path}")
            with open(file_path, 'rb') as file:
                response = HttpResponse(file.read(), content_type=content_type)
                response['Content-Disposition'] = f'attachment; filename="{os.path.basename(file_path)}"'
                return response
        else:
            logging.warning("No file available for download.")
            return HttpResponse("No file available for download.")
    except Exception as e:
        logging.error(f"An error occurred while downloading the file: {e}")
        return HttpResponse(f"An error occurred while downloading the file: {e}")



